import os
import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

def inicializar_livro_master():
    doc_master = docx.Document()
    for section in doc_master.sections:
        section.top_margin = Inches(1.18)    
        section.left_margin = Inches(1.18)   
        section.bottom_margin = Inches(0.78) 
        section.right_margin = Inches(0.78)  

    style = doc_master.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    return doc_master

def adicionar_capa_master(doc):
    for _ in range(3):
        doc.add_paragraph()
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run("RENÊ APARECIDO BUENO")
    run_author.bold = True
    run_author.font.size = Pt(14)

    for _ in range(6):
        doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(
        "TRATADO DE MUSICOLOGIA E HARMONIA TONAL:\n"
        "DA EPISTEMOLOGIA ACÚSTICA ÀS PRÁTICAS INSTRUMENTAIS\n\n"
        "COMPÊNDIO COMPLETO COM APÊNDICE DE MAPEAMENTO INSTRUMENTAL"
    )
    run_title.bold = True
    run_title.font.size = Pt(16)

    for _ in range(10):
        doc.add_paragraph()
    p_local = doc.add_paragraph()
    p_local.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_local = p_local.add_run("MARINGÁ - PR\n2026")
    run_local.bold = True
    run_local.font.size = Pt(12)
    doc.add_page_break()

def clonar_conteudo_arquivo(doc_master, caminho_arquivo):
    if not os.path.exists(caminho_arquivo):
        return
    doc_origem = docx.Document(caminho_arquivo)
    for elemento in doc_origem.element.body:
        if elemento.tag.endswith('p'):
            p_origem = docx.text.paragraph.Paragraph(elemento, doc_origem)
            if p_origem.text.strip():
                p_novo = doc_master.add_paragraph()
                p_novo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_novo.paragraph_format.line_spacing = 1.5
                p_novo.paragraph_format.space_after = Pt(12)
                if p_origem.text.startswith(("CAPÍTULO", "1", "2", "3", "4", "5", "6", "SUMÁRIO", "REFERÊNCIAS")):
                    p_novo.paragraph_format.first_line_indent = Inches(0)
                    p_novo.add_run(p_origem.text).bold = True
                else:
                    p_novo.paragraph_format.first_line_indent = Inches(0.49)
                    p_novo.add_run(p_origem.text)
        elif elemento.tag.endswith('tbl'):
            tbl_origem = docx.table.Table(elemento, doc_origem)
            tbl_nova = doc_master.add_table(rows=len(tbl_origem.rows), cols=len(tbl_origem.columns))
            tbl_nova.style = 'Table Grid'
            for r_idx, row in enumerate(tbl_origem.rows):
                for c_idx, cell in enumerate(row.cells):
                    tbl_nova.rows[r_idx].cells[c_idx].text = cell.text
                    if r_idx == 0:
                        tbl_nova.rows[r_idx].cells[c_idx].paragraphs[0].runs[0].bold = True

def gerar_diagrama_vetorial(nome_arquivo, nome_acorde, dedos, pasta_imagens):
    """Desenha matematicamente o diagrama do acorde e salva em alta resolução (PNG)."""
    caminho_completo = os.path.join(pasta_imagens, nome_arquivo)
    
    # Configuração da grade do braço do violão (6 cordas verticais, 5 casas horizontais)
    fig, ax = plt.subplots(figsize=(3.5, 4))
    
    # Desenha as 5 linhas horizontais (Casas 1 a 5)
    for i in range(1, 6):
        ax.axhline(i, color='black', linewidth=1.5, zorder=1)
    
    # Desenha as 6 linhas verticais (Cordas do Violão)
    for i in range(1, 7):
        ax.axvline(i, color='gray', linewidth=1.2, zorder=1)
        
    # Destaca a pestana superior (Nut) com uma linha preta grossa
    ax.axvline(1, color='black', linewidth=5, zorder=2)
    
    # Adiciona as marcações dos dedos baseadas na matriz
    # Formato: (Corda de 1 a 6 de baixo para cima, Casa, Número do Dedo)
    for corda, casa, dedo in dedos:
        if casa > 0:
            # Plota a bolinha preta do dedo no cruzamento exato da corda e da casa
            ax.scatter(casa + 0.5, corda, color='black', s=450, zorder=3)
            # Adiciona o número do dedo em branco dentro da bolinha
            ax.text(casa + 0.5, corda, str(dedo), color='white', ha='center', va='center', 
                    fontname='Arial', fontsize=11, fontweight='bold', zorder=4)

    # Ajustes estéticos finais do gráfico para parecer um livro de música real
    ax.set_title(nome_acorde, fontname='Times New Roman', fontsize=14, fontweight='bold', pad=15)
    ax.set_xlim(0.8, 5.2)
    ax.set_ylim(0.5, 6.5)
    
    # Nomeação das cordas na lateral esquerda do diagrama (1=E aguda, 6=E grave)
    ax.set_yticks([1, 2, 3, 4, 5, 6])
    ax.set_yticklabels(['E (1ª)', 'B (2ª)', 'G (3ª)', 'D (4ª)', 'A (5ª)', 'E (6ª)'], fontname='Times New Roman', fontsize=10)
    
    # Nomeação das casas na parte inferior
    ax.set_xticks([1.5, 2.5, 3.5, 4.5])
    ax.set_xticklabels(['Casa 1', 'Casa 2', 'Casa 3', 'Casa 4'], fontname='Times New Roman', fontsize=10)
    
    ax.grid(False)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['left'].set_visible(False)
    
    # Inverte o eixo y para que a 1ª corda (E aguda) fique em cima
    ax.invert_yaxis()
    
    # Salva a imagem sem bordas brancas sobressalentes
    plt.tight_layout()
    plt.savefig(caminho_completo, dpi=200)
    plt.close()
    return caminho_completo

def inserir_diagrama_no_documento(doc, nome_acorde, caminho_imagem):
    """Insere o diagrama gerado centralizado e formatado segundo padrões editoriais."""
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(14)
    p_titulo.paragraph_format.keep_with_next = True
    p_titulo.add_run(f"Diagrama Ilustrativo: {nome_acorde}").bold = True
    
    if caminho_imagem and os.path.exists(caminho_imagem):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(4)
        p_img.add_run().add_picture(caminho_imagem, width=Cm(5.0))
        
        p_legenda = doc.add_paragraph()
        p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_legenda.paragraph_format.space_after = Pt(16)
        run_leg = p_legenda.add_run(f"Figura: Mapeamento posicional para o acorde de {nome_acorde}.")
        run_leg.font.size = Pt(10)
        run_leg.italic = True


# --- CONFIGURAÇÃO DE DIRETÓRIOS ---
diretorio_base = r"C:\Users\rene.bueno\Desktop\livro"
pasta_imagens = os.path.join(diretorio_base, "imagens_acordes")

if not os.path.exists(pasta_imagens):
    os.makedirs(pasta_imagens)

livro_master = inicializar_livro_master()
adicionar_capa_master(livro_master)

# Unificação das partes de texto
arquivos_docs = [
    "Tratado_de_Musicologia_Parte1.docx", "Tratado_de_Musicologia_Parte2.docx",
    "Tratado_de_Musicologia_Parte3.docx", "Tratado_de_Musicologia_Parte4.docx",
    "Tratado_de_Musicologia_Parte5.docx", "Tratado_de_Musicologia_Parte6.docx"
]

for nome_arq in arquivos_docs:
    caminho = os.path.join(diretorio_base, nome_arq)
    if os.path.exists(caminho):
        clonar_conteudo_arquivo(livro_master, caminho)
        livro_master.add_page_break()


# =========================================================================
# APÊNDICE B: CENTRAL DE ACORDES (TÉTRADES, TENSÕES E INVERSÕES)
# =========================================================================
p_apendice = livro_master.add_paragraph()  # <--- CORRIGIDO AQUI!
run_ap = p_apendice.add_run("APÊNDICE B — DICIONÁRIO GEOMÉTRICO DE ACORDES (REPOSITÓRIO VISUAL EXPANDIDO)")
run_ap.bold = True
run_ap.font.size = Pt(14)
p_apendice.paragraph_format.space_before = Pt(24)

# Mapeamento técnico atualizado (Corda de 1 a 6 de cima para baixo no gráfico, Casa, Número do Dedo)
banco_acordes_vetoriais = [
    # --- TRÍADES BÁSICAS ---
    ("Dó Maior (C)", "C_maior.png", [(5, 3, 3), (4, 2, 2), (2, 1, 1)]),
    ("Lá Maior (A)", "A_maior.png", [(4, 2, 1), (3, 2, 2), (2, 2, 3)]),
    ("Sol Maior (G)", "G_maior.png", [(6, 3, 3), (5, 2, 2), (1, 3, 4)]),
    ("Mi Maior (E)", "E_maior.png", [(5, 2, 2), (4, 2, 3), (3, 1, 1)]),
    ("Ré Maior (D)", "D_maior.png", [(3, 2, 1), (2, 3, 3), (1, 2, 2)]),
    ("Lá Menor (Am)", "A_menor.png", [(4, 2, 2), (3, 2, 3), (2, 1, 1)]),
    ("Mi Menor (Em)", "E_menor.png", [(5, 2, 2), (4, 2, 3)]),
    ("Ré Menor (Dm)", "D_menor.png", [(3, 2, 2), (2, 3, 3), (1, 1, 1)]),

    # --- COM SÉTIMA MAIOR (C7M, G7M, A7M, D7M) ---
    ("Dó com Sétima Maior (C7M)", "C7M.png", [(5, 3, 3), (4, 2, 2), (2, 0, 0)]),
    ("Sol com Sétima Maior (G7M)", "G7M.png", [(6, 3, 2), (5, 2, 1), (1, 2, 3)]),
    ("Lá com Sétima Maior (A7M)", "A7M.png", [(4, 2, 1), (3, 1, 2), (2, 2, 3)]),
    ("Ré com Sétima Maior (D7M)", "D7M.png", [(3, 2, 1), (2, 2, 1), (1, 2, 1)]),

    # --- COM SÉTIMA DOMINANTE/MENOR (C7, A7, E7, D7) ---
    ("Dó com Sétima (C7)", "C7.png", [(5, 3, 3), (4, 2, 2), (3, 3, 4), (2, 1, 1)]),
    ("Lá com Sétima (A7)", "A7.png", [(4, 2, 1), (2, 2, 2)]),
    ("Mi com Sétima (E7)", "E7.png", [(5, 2, 2), (3, 1, 1)]),
    ("Ré com Sétima (D7)", "D7.png", [(3, 2, 2), (2, 1, 1), (1, 2, 3)]),

    # --- COM NONA (C9, G9, D9, Am9) ---
    ("Dó com Nona (C9)", "C9.png", [(5, 3, 2), (4, 2, 1), (2, 3, 3), (1, 3, 4)]),
    ("Sol com Nona (G9)", "G9.png", [(6, 3, 2), (5, 2, 1), (2, 3, 3), (1, 3, 4)]),
    ("Ré com Nona (D9)", "D9.png", [(3, 2, 1), (2, 3, 3)]),
    ("Lá Menor com Nona (Am9)", "Am9.png", [(4, 2, 2), (3, 5, 4), (2, 0, 0)]),

    # --- DIMINUTOS (C° / Cdim, G° / Gdim) ---
    ("Dó Diminuto (Cdim)", "C_dim.png", [(5, 3, 3), (4, 4, 1), (3, 5, 4), (2, 4, 2)]),
    ("Sol Diminuto (Gdim)", "G_dim.png", [(4, 2, 1), (3, 3, 3), (2, 2, 2), (1, 3, 4)]),

    # --- NOTAS COM BAIXO ALTERADO / INVERSÕES (C/E, G/B, D/F#) ---
    ("Dó com Baixo em Mi (C/E)", "C_sob_E.png", [(6, 0, 0), (5, 3, 3), (4, 2, 2), (2, 1, 1)]),
    ("Sol com Baixo em Sí (G/B)", "G_sob_B.png", [(5, 2, 1), (1, 3, 2)]),
    ("Ré com Baixo em Fá Sustenido (D/F#)", "D_sob_Fsharp.png", [(6, 2, 1), (3, 2, 2), (2, 3, 4), (1, 2, 3)]),
]

print("Iniciando mecanismo de renderização gráfica interna...")

for nome, arquivo_local, mapeamento_dedos in banco_acordes_vetoriais:
    print(f"-> Desenhando e salvando em disco: {nome}")
    caminho_da_imagem = gerar_diagrama_vetorial(arquivo_local, nome, mapeamento_dedos, pasta_imagens)
    inserir_diagrama_no_documento(livro_master, nome, caminho_da_imagem)

# --- SALVAMENTO FINAL ---
caminho_salvamento = os.path.join(diretorio_base, "Tratado_de_Musicologia_Compendio_Completo.docx")
livro_master.save(caminho_salvamento)

print("\n==========================================================")
print("PROCESSO EXECUTADO COM SUCESSO ABSOLUTO!")
print(f"Os novos diagramas (Tétrades e Inversões) estão em: {pasta_imagens}")
print(f"Compêndio completo estruturado com sucesso em: {caminho_salvamento}")
print("==========================================================")